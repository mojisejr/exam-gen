"""
Test Services
Tests the Muscle: Document generation logic.
"""
import os
from docx import Document
from server.services.doc_generator import generate_docx


def test_docx_structure(mock_worksheet, tmp_path):
    """Test that generated DOCX has correct structure."""
    output_path = tmp_path / "test_output.docx"
    
    # Generate document
    generate_docx(mock_worksheet, str(output_path))
    
    # Verify file exists
    assert output_path.exists()
    
    # Read document back
    doc = Document(str(output_path))
    
    # Verify content exists
    text = "\n".join([para.text for para in doc.paragraphs])
    
    assert "Test Worksheet" in text  # Title
    assert "Mathematics" in text  # Subject
    assert "What is 2+2?" in text  # Question 1
    assert "What is the capital of France?" in text  # Question 2


def test_docx_answer_key(mock_worksheet, tmp_path):
    """Test that answer key table is present."""
    output_path = tmp_path / "test_answer_key.docx"
    
    generate_docx(mock_worksheet, str(output_path))
    
    doc = Document(str(output_path))
    
    # Check for tables (answer key)
    assert len(doc.tables) > 0
    
    # Verify table structure
    table = doc.tables[0]
    assert len(table.rows) >= 2  # Header + at least one answer row


def test_docx_metadata(mock_worksheet, tmp_path):
    """Test that document metadata is set correctly."""
    output_path = tmp_path / "test_metadata.docx"
    
    generate_docx(mock_worksheet, str(output_path))
    
    doc = Document(str(output_path))
    
    assert doc.core_properties.title == "Test Worksheet"
    assert doc.core_properties.subject == "Mathematics"


def test_docx_true_false_layout(mock_true_false_worksheet, tmp_path):
    """Test that true/false layout renders checkbox lines."""
    output_path = tmp_path / "test_true_false.docx"

    generate_docx(mock_true_false_worksheet, str(output_path))

    doc = Document(str(output_path))
    text = "\n".join([para.text for para in doc.paragraphs])

    assert "( ) ถูก" in text
    assert "( ) ผิด" in text


def test_docx_subjective_layout(mock_subjective_worksheet, tmp_path):
    """Test that subjective layout renders blank lines for answers."""
    output_path = tmp_path / "test_subjective.docx"

    generate_docx(mock_subjective_worksheet, str(output_path))

    doc = Document(str(output_path))
    text = "\n".join([para.text for para in doc.paragraphs])

    assert "........................................................" in text


def test_docx_thai_stress_text(tmp_path):
    """Test Thai complex text rendering in DOCX."""
    output_path = tmp_path / "test_thai_text.docx"

    worksheet = {
        "title": "แบบทดสอบภาษาไทย",
        "subject": "ภาษาไทย",
        "target_level": "มัธยมศึกษาปีที่ 3",
        "items": [
            {
                "id": 1,
                "question": "อธิบายคำว่า 'เกลี้ยงเกลา' และยกตัวอย่างจากวรรณคดีไทย",
                "type": "subjective",
                "options": None,
                "correct_answer": "คำอธิบายเชิงวิเคราะห์",
                "explanation": None,
                "image_prompt": None,
            }
        ],
    }

    from server.schemas import Worksheet

    generate_docx(Worksheet(**worksheet), str(output_path))

    doc = Document(str(output_path))
    text = "\n".join([para.text for para in doc.paragraphs])

    assert "แบบทดสอบภาษาไทย" in text
    assert "เกลี้ยงเกลา" in text
