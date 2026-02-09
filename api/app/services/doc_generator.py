"""
Document Generator Service
Handles .docx file generation with proper Thai language formatting.
"""
from docx import Document
from docx.shared import Pt, Inches

from app.schemas import QuestionType, Worksheet


def _add_multiple_choice(doc: Document, item) -> None:
    options = item.options or []
    for opt in options:
        p_opt = doc.add_paragraph()
        p_opt.paragraph_format.left_indent = Inches(0.5)
        text = opt.text if opt.text.strip().startswith(opt.label) else f"{opt.label}) {opt.text}"
        p_opt.add_run(text)


def _add_true_false(doc: Document) -> None:
    for label in ("ถูก", "ผิด"):
        p_opt = doc.add_paragraph()
        p_opt.paragraph_format.left_indent = Inches(0.5)
        p_opt.add_run(f"( ) {label}")


def _add_subjective(doc: Document) -> None:
    p_line = doc.add_paragraph()
    p_line.paragraph_format.left_indent = Inches(0.5)
    p_line.add_run("........................................................")


def generate_docx(worksheet: Worksheet, output_path: str) -> None:
    """
    Generates a .docx file from a Worksheet object.
    
    Args:
        worksheet: Worksheet object containing exam data.
        output_path: Absolute path where the .docx file should be saved.
    
    Raises:
        Exception: If file generation fails.
    """
    print(f"\n[Publisher] Generating DOCX at {output_path}...")
    doc = Document()
    
    # Set Metadata
    core_properties = doc.core_properties
    core_properties.title = worksheet.title
    core_properties.subject = worksheet.subject

    # Style Adjustments (Thai-friendly font sizing)
    style = doc.styles['Normal']
    style.font.size = Pt(14)
    
    # Title Section
    heading = doc.add_heading(worksheet.title, 0)
    heading.alignment = 1  # Center
    
    p_info = doc.add_paragraph()
    p_info.alignment = 1  # Center
    p_info.add_run(f"วิชา: {worksheet.subject}").bold = True
    p_info.add_run(f" | ระดับชั้น: {worksheet.target_level}")
    
    doc.add_paragraph("_" * 50).alignment = 1  # Divider

    # Questions Section
    for item in worksheet.items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run(f"{item.id}. {item.question}")
        run.bold = True
        run.font.size = Pt(14)
        
        # Image Placeholder (if needed)
        if item.image_prompt:
            p_img = doc.add_paragraph()
            p_img.alignment = 1  # Center
            run_img = p_img.add_run(f"[🖼️ INSERT IMAGE HERE: {item.image_prompt}]")
            run_img.font.color.rgb = None 
            run_img.italic = True
            run_img.font.size = Pt(10)
            doc.add_paragraph()  # Add spacing
        
        # Options / Input Area
        if item.type == QuestionType.TRUE_FALSE:
            _add_true_false(doc)
        elif item.type == QuestionType.SUBJECTIVE:
            _add_subjective(doc)
        else:
            _add_multiple_choice(doc, item)
        
    # Answer Key (New Page)
    doc.add_page_break()
    doc.add_heading("เฉลยคำตอบ / Answer Key", level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ข้อ (No.)'
    hdr_cells[1].text = 'คำตอบ (Ans)'
    hdr_cells[2].text = 'คำอธิบาย (Explanation)'

    for item in worksheet.items:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.id)
        row_cells[1].text = item.correct_answer
        row_cells[2].text = item.explanation or "-"

    doc.save(output_path)
    print("Done! File saved.")
